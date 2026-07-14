"""Download Telegram documents and extract text for Cursor prompts."""

from __future__ import annotations

import logging
from io import BytesIO

from telegram import Message

logger = logging.getLogger(__name__)

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_SUPPORTED_MIMES = frozenset({PDF_MIME, DOCX_MIME})


class UnsupportedDocumentError(Exception):
    """Raised when a message has no supported document."""


class DocumentTooLargeError(Exception):
    """Raised when the downloaded file exceeds the configured byte limit."""


class DocumentEmptyError(Exception):
    """Raised when text could not be extracted (e.g. scanned PDF)."""


def _normalize_mime(mime_type: str | None) -> str:
    return (mime_type or "").strip().lower()


def _guess_kind(mime_type: str | None, filename: str | None) -> str | None:
    mime = _normalize_mime(mime_type)
    if mime in _SUPPORTED_MIMES:
        return "pdf" if mime == PDF_MIME else "docx"
    name = (filename or "").strip().lower()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    return None


def is_supported_document_message(message: Message) -> bool:
    doc = message.document
    if doc is None:
        return False
    return _guess_kind(doc.mime_type, doc.file_name) is not None


def document_default_prompt() -> str:
    return "Проанализируй вложенный документ и ответь по сути."


def build_prompt_text(caption: str | None) -> str:
    text = (caption or "").strip()
    return text if text else document_default_prompt()


def truncate_document_text(text: str, *, max_chars: int) -> tuple[str, bool]:
    cleaned = text.strip()
    if len(cleaned) <= max_chars:
        return cleaned, False
    return f"{cleaned[:max_chars]}\n\n[... текст обрезан ...]", True


def build_document_prompt(
    *,
    caption: str | None,
    filename: str,
    text: str,
) -> str:
    user_part = build_prompt_text(caption)
    return (
        f"--- Вложение: {filename} ---\n"
        f"{text}\n"
        f"---\n\n"
        f"{user_part}"
    )


def extract_text_from_bytes(data: bytes, *, kind: str) -> str:
    if kind == "pdf":
        return _extract_pdf_text(data)
    if kind == "docx":
        return _extract_docx_text(data)
    raise UnsupportedDocumentError(f"unsupported kind: {kind}")


def _extract_pdf_text(data: bytes) -> str:
    import fitz

    with fitz.open(stream=data, filetype="pdf") as pdf:
        parts: list[str] = []
        for page in pdf:
            page_text = page.get_text().strip()
            if page_text:
                parts.append(page_text)
    return "\n\n".join(parts).strip()


def _extract_docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip()
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


from beachops.services.telegram_download import (
    TelegramFileDownloadError,
    download_telegram_file_bytes,
)


async def download_telegram_document(
    message: Message,
    *,
    max_bytes: int,
    retries: int | None = None,
    retry_delay_sec: float | None = None,
) -> tuple[bytes, str, str]:
    """Download a PDF/DOCX from Telegram. Returns (bytes, kind, filename)."""
    doc = message.document
    if doc is None:
        raise UnsupportedDocumentError("no document in message")

    kind = _guess_kind(doc.mime_type, doc.file_name)
    if kind is None:
        raise UnsupportedDocumentError(
            f"unsupported mime: {doc.mime_type!r}, file: {doc.file_name!r}"
        )

    if doc.file_size is not None and doc.file_size > max_bytes:
        raise DocumentTooLargeError(doc.file_size)

    try:
        data = await download_telegram_file_bytes(
            message.get_bot(),
            doc.file_id,
            retries=retries,
            retry_delay_sec=retry_delay_sec,
        )
    except TelegramFileDownloadError as exc:
        raise DocumentEmptyError("telegram document download failed") from exc

    if not data:
        raise DocumentEmptyError("empty download")
    if len(data) > max_bytes:
        raise DocumentTooLargeError(len(data))

    filename = doc.file_name or f"document.{kind}"
    return data, kind, filename


async def extract_message_document_text(
    message: Message,
    *,
    max_bytes: int,
    max_chars: int,
) -> tuple[str, str, bool]:
    """Download and extract text. Returns (filename, text, was_truncated)."""
    data, kind, filename = await download_telegram_document(message, max_bytes=max_bytes)
    text = extract_text_from_bytes(data, kind=kind)
    if not text:
        raise DocumentEmptyError(f"no extractable text in {filename}")

    truncated_text, was_truncated = truncate_document_text(text, max_chars=max_chars)
    return filename, truncated_text, was_truncated
